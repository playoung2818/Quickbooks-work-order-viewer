import os
import json
import logging
import pdfplumber
import pandas as pd
from docx import Document
from models import db, PDFFileLog, WordFileLog

def validate_paths(paths):
    for path in paths:
        if not os.path.exists(path):
            logging.error(f"Path does not exist: {path}")
        else:
            logging.info(f"Valid path: {path}")

def process_pdf_files(folder_path):
    pdf_data = {}
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.pdf'):
                file_path = os.path.join(root, file)
                order_id = os.path.splitext(file)[0]  
                logging.info(f"Processing PDF file: {file}")

                # Store only metadata, no table extraction
                pdf_data[order_id] = {
                    "file_name": file,
                    "file_path": file_path
                }
    return pdf_data

def process_word_files(folder_path):
    word_data = {}
    if not os.path.exists(folder_path):
        logging.error(f"Folder path does not exist: {folder_path}")
        return word_data

    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.docx'):
                try:
                    file_path = os.path.join(root, file)
                    order_id = os.path.splitext(file)[0]
                    logging.info(f"Processing Word file: {file}")
                    product_details = extract_product_details_from_word(file_path)
                    word_data[order_id] = {
                        "file_name": file,
                        "file_path": file_path,
                        "product_details": product_details
                    }
                except Exception as e:
                    logging.error(f"Error processing Word file {file}: {e}")
    return word_data

def extract_product_details_from_word(file_path):
    try:
        if not os.path.exists(file_path):
            logging.warning(f"Word file path does not exist: {file_path}")
            return []

        document = Document(file_path)
        if not document.tables:
            logging.warning(f"No tables found in Word file: {file_path}")
            return []

        table = document.tables[0]
        product_details = []
        for i, row in enumerate(table.rows[1:]):  # Skip header row
            cells = row.cells
            if len(cells) < 4:
                logging.warning(f"Row {i + 1} in {file_path} has insufficient cells.")
                continue

            product_details.append({
                "product_number": cells[0].text.strip(),
                "qty": cells[1].text.strip(),
                "sn": cells[2].text.strip(),
                "notes": cells[3].text.strip()
            })
        return product_details
    except Exception as e:
        logging.error(f"Error processing Word file {file_path}: {e}")
        return []

def load_pdf_files_to_db(pdf_data):
    for order_id, data in pdf_data.items():
        try:
            existing_entry = PDFFileLog.query.filter_by(order_id=order_id, file_name=data['file_name']).first()
            if existing_entry:
                logging.info(f"Duplicate PDF entry detected for Order ID {order_id}, File: {data['file_name']}. Skipping insert.")
                continue

            logging.info(f"Inserting into DB: Order ID: {order_id}, File Name: {data['file_name']}")

            db.session.add(PDFFileLog(
                order_id=order_id,
                file_name=data['file_name'],
                file_path=data['file_path']
            ))
        except Exception as e:
            db.session.rollback()
            logging.error(f"Database insert error for order ID {order_id}: {e}")

    db.session.commit()
    logging.info("All PDF file data loaded into the database.")

def load_word_files_to_db(word_data):
    for order_id, data in word_data.items():
        try:
            existing_entry = WordFileLog.query.filter_by(order_id=order_id, file_name=data['file_name']).first()
            if existing_entry:
                logging.info(f"Duplicate Word entry detected for Order ID {order_id}. Skipping.")
                continue

            db.session.add(WordFileLog(
                order_id=order_id,
                file_name=data['file_name'],
                product_details=data['product_details'],
                file_path=data['file_path']
            ))
        except Exception as e:
            logging.error(f"Error saving Word file data for order ID {order_id}: {e}")
    db.session.commit()
    logging.info("Word files successfully saved to database.")

def process_all_work_order_pdfs():
    WO_PDF_FOLDER = r"\\Quickbook2024\d\Drive D\QuickBooks\2- Year 2024\Work Order- WO"
    WO_PDF_FOLDER2 = r"\\Quickbook2024\d\Drive D\QuickBooks\3- Year 2025\Work Order- WO"
    logging.info("Processing all Work Order PDF files...")
    pdf_data = process_pdf_files(WO_PDF_FOLDER)
    pdf_data.update(process_pdf_files(WO_PDF_FOLDER2))
    if pdf_data:
        load_pdf_files_to_db(pdf_data)
    else:
        logging.info("No PDF files found in the Work Order folder.")
    logging.info("All Work Order PDF files processed and saved.")

def process_all_work_order_words():
    WO_WORD_FOLDER = r"C:\Users\Admin\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2024"
    WO_WORD_FOLDER2 = r"C:\Users\Admin\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2025"
    logging.info("Processing all Work Order Word files...")
    word_data = process_word_files(WO_WORD_FOLDER)
    word_data.update(process_word_files(WO_WORD_FOLDER2))
    if word_data:
        load_word_files_to_db(word_data)
    else:
        logging.info("No Word files found in the Work Order folder.")
    logging.info("All Work Order Word files processed and saved.")

def update_pdf_file_paths(old_base_path, new_base_path):
    entries = PDFFileLog.query.all()
    updated_count = 0

    for entry in entries:
        if entry.file_path.startswith(old_base_path):
            entry.file_path = entry.file_path.replace(old_base_path, new_base_path)
            updated_count += 1

    db.session.commit()
    logging.info(f"✅ Updated file paths for {updated_count} PDF entries.")
